import collections

import docx
from docx.enum.text import WD_COLOR_INDEX

from . import style


def correcting(styles, data):
    if not styles:
        return data

    elif isinstance(s := styles[-1], style.InPlaceStyle):
        data = (
            correcting(styles[:-1], data[: s.start])
            + wrapping([styles[-1]], data[s.start : s.end])
            + data[s.end :]
        )

    return data


def tag(s):
    return s[0]


def attributes(s):
    attr = " ".join([f'{k}="{v}"' for k, v in s[1].items()])
    if attr:
        return " " + attr
    return ""


def make_wrap(s):
    if not isinstance(s, style.Style):
        return ("<" + tag(s) + attributes(s) + ">", "</" + tag(s) + ">")

    if isinstance(s, style.MarginTop):
        return f'<div style="margin-top: {s.lines}em;">', "</div>"

    if isinstance(s, style.Highlight):
        return f'<span style="background-color: {s.color}">', "</span>"

    if isinstance(s, style.Hide):
        return f'<div style="display: none">', "</div>"

    if isinstance(s, style.Repetition):
        label = (
            f'<a href="#Label{s.label}" title="{s.i + 1} повтор из {s.n}">'
            '<sup class="badge badge-secondary" style="font-size:10px;">'
            f"{s.label} ({s.i + 1}/{s.n})"
            "</sup>"
            "</a>"
        )
        if s.i == 0:
            return f'<span id="Label{s.label}">', f"</span>{label}"

        return "", f"{label}"

    if isinstance(s, style.FontColor):
        return f'<span style="color: {s.color.value};">', "</span>"

    return "", ""


def wrapping(styles, data):
    wrap = list(zip(*map(make_wrap, styles)))
    if len(wrap) > 0:
        return "\n".join(wrap[0]) + data + "\n".join(wrap[1][-1::-1])
    else:
        return data


def is_paragraph_id(id):
    return id != "origin" and len(id.split("s")) == 1 and len(id.split("p")) == 2


def assemble(struct, styles, data, out_type="html"):
    def html_compiler(node):
        """ """
        if isinstance(node, str):
            return correcting(styles.get(node, []), data[node[1:]])

        id, next_nodes = node

        if is_paragraph_id(id):
            return (
                "<p>"
                + wrapping(styles.get(id, []), " ".join(map(html_compiler, next_nodes)))
                + "</p>"
            )
        return wrapping(styles.get(id, []), "\n".join(map(html_compiler, next_nodes)))

    if out_type == "html":
        return html_compiler(struct)

    def compile_txt(node):
        node_id, subnodes = node

        if len(subnodes) == 1 and isinstance(leaf := subnodes[0], str):
            leaf_id = leaf[1:]
            indent = " "
            prefix = ""
            sentence = data[leaf_id]

            for s in styles[node_id]:
                if isinstance(s, style.ParagraphIndent):
                    indent = "\n" * s.width

                if isinstance(s, style.Repetition):
                    prefix = f"[{s.label} | {s.i + 1}/{s.n}]" + prefix

                    if s.i > 0:
                        limit = 80
                        if len(sentence) > limit:
                            sentence = sentence[:limit].rstrip() + "....."

                if isinstance(s, style.Exclusive):
                    prefix = "[!]" + prefix

                if isinstance(s, style.Hide):
                    sentence = ""

            return f"{indent}{prefix}{sentence}"

        content = "".join(map(compile_txt, subnodes))

        if is_paragraph_id(node_id):
            n_breaklines = 2
            end = ""
            for s in styles[node_id]:
                if isinstance(s, style.MarginTop):
                    n_breaklines += s.lines
                elif isinstance(s, style.FirstParagraph):
                    n_breaklines -= 2
                elif isinstance(s, style.FinalNewline):
                    end = "\n"
                if isinstance(s, style.Hide):
                    return ""

            return n_breaklines * "\n" + content + end

        return content

    if out_type == "txt":
        return compile_txt(struct)

    if out_type == "docx":
        return assemble_docx(struct, styles, data)


def assemble_docx(struct, styles, data, parent=None):
    node_id, subnodes = struct

    if len(subnodes) == 1 and isinstance(leaf := subnodes[0], str):
        leaf_id = leaf[1:]
        parent.add_run(" ")
        sentence = data[leaf_id]
        for s in styles[node_id]:
            if isinstance(s, style.Repetition):
                parent.add_run(f"[{s.label} | {s.i + 1}/{s.n}]")

            if isinstance(s, style.Exclusive):
                parent.add_run("[!]")

            if isinstance(s, style.Hide):
                return parent

        inplace_styles = [s for s in styles[leaf] if isinstance(s, style.InPlaceStyle)]
        if len(inplace_styles) == 0:
            parent.add_run(sentence)
        else:
            cursor = 0
            for s in inplace_styles:
                parent.add_run(sentence[cursor : s.start])
                text = sentence[s.start : s.end]
                r = parent.add_run(text)
                if isinstance(s, style.Highlight):
                    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

                cursor = s.end
            parent.add_run(sentence[cursor:])
        return parent

    if is_paragraph_id(node_id):
        for s in styles[node_id]:
            if isinstance(s, style.Hide):
                return parent

            if isinstance(s, style.MarginTop):
                for i in range(s.lines):
                    parent.add_paragraph()

        p = parent.add_paragraph()
        for subnode in subnodes:
            assemble_docx(subnode, styles, data, parent=p)
        return parent

    parent = parent or docx.Document()
    for subnode in subnodes:
        assemble_docx(subnode, styles, data, parent=parent)

    return parent
