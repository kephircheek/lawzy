import collections
import itertools
import json
import os
from datetime import datetime
from uuid import uuid4

import docx
from flask import (
    Blueprint,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)

from lawzy import core
from lawzy.app.models import (
    Data,
    KeywordEntries,
    Struct,
    Style,
    document_ids,
    document_name,
)
from lawzy.config import UPLOAD_FOLDER

INDENT = 2
HIDE_PARS_WITHOUT_KEYWORDS = "hide_pars_without_keywords"

# Define the blueprint: 'auth', set its url prefix: app.url/auth
aggregator = Blueprint("aggregator", __name__, url_prefix="/aggregator")

# Set the route and accepted methods
@aggregator.route("/uploader", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":

        token = uuid4()
        session["token"] = str(token)
        session["toggle:reduce"] = False
        session["reduced"] = False
        os.makedirs(f"{UPLOAD_FOLDER}/{token}")

        project_config = {
            "DATE_CREATED": str(datetime.now()),
            "SPLIT_SENTENCE_PATTERN": r"(?<=[^( г| N| ном)]\.)\s+(?=[^\Wа-яa-z0-9])",
            "SPLIT_CASE_PATTERN": "Документ предоставлен КонсультантПлюс",
            "CASE_NUMBER_PATTERN": r"(?<=по делу\s)[^\n]+",
        }

        with open(f"{UPLOAD_FOLDER}/{token}/config.json", "w") as f:
            json.dump(project_config, f, indent=INDENT)

        for document_id, f in enumerate(request.files.getlist("file")):
            filename, extention = os.path.splitext(os.path.basename(f.filename))
            path_source = f"{UPLOAD_FOLDER}/{token}/source{extention}"
            f.save(path_source)

            with open(path_source, "rb") as f:
                if extention == ".docx":
                    doc = docx.Document(f)
                    struct, styles, data = core.parser.parse_docx(
                        doc, project_config["SPLIT_SENTENCE_PATTERN"]
                    )
                else:
                    text = f.read().decode("utf-8")
                    struct, styles, data = core.parser.parse_txt(
                        text, project_config["SPLIT_SENTENCE_PATTERN"]
                    )

            document_config = {
                "FILENAME": filename,
                "EXTENSION": extention,
                "PATH_SOURCE": path_source,
            }
            path_document = UPLOAD_FOLDER / f"{token}/{document_id}"
            path_document.mkdir()
            with open(path_document / "config.json", "w") as f:
                json.dump(document_config, f, indent=INDENT)

            Struct(token, document_id).post(struct)
            Style(token, document_id).post(styles)
            Data(token, document_id).post({id: [item] for id, item in data.items()})
            KeywordEntries(token, document_id).post(dict())

        return redirect(url_for("aggregator.document", document_id=str(0)))


@aggregator.route("document/<document_id>", methods=["GET", "POST"])
def document(document_id: str):
    if "token" not in session:
        return redirect(url_for("hello"))

    token = session["token"]
    session["document_id"] = document_id

    if "toggle:reduce" not in session:
        session["toggle:reduce"] = False

    if HIDE_PARS_WITHOUT_KEYWORDS not in session:
        session[HIDE_PARS_WITHOUT_KEYWORDS] = False

    checked = "checked" if session["toggle:reduce"] else ""

    struct = Struct(token, document_id).get()
    styles = Style(token, document_id).get()
    data = Data(token, document_id).sentences
    profit = ""
    if session["toggle:reduce"] == True:
        labels = Data(token, document_id).labels
        number_all = len(labels)
        number_of_group = len(set(labels.values())) - 1
        number_non_grouped = len(list(filter(lambda x: x == -1, labels.values())))
        remain = (number_non_grouped + number_of_group) / number_all
        profit = "-%s%%" % round(100 * (1 - remain))

    else:
        labels = None
    keywords = KeywordEntries(token, document_id).get().keys()
    content = core.compiler.assemble(
        struct, styles, data, labels=labels, mute=session["toggle:reduce"], limit=80
    )
    return render_template(
        "aggregator/document.html",
        content=content,
        keywords=keywords,
        checked=checked,
        profit=profit,
        reduced=session["reduced"],
        session=session,
        documents=sorted((id, document_name(token, id)) for id in document_ids(token)),
    )


@aggregator.route("/reduce", methods=["GET"])
def reduce():
    token = session["token"]

    for document_id in document_ids(token):
        if not session["reduced"]:
            labels = core.group(Data(token, document_id).sentences)
            Data(token, document_id).add(labels)
            dubs = core.dublicates(labels)
            Data(token, document_id).add(dubs)

    session["reduced"] = True
    session["toggle:reduce"] = not session["toggle:reduce"]

    return redirect(url_for("aggregator.document", document_id=session["document_id"]))


@aggregator.route("/rating", methods=["GET"])
def rating():

    token = session["token"]
    document_id = session["document_id"]
    path = os.path.abspath(f"{UPLOAD_FOLDER}/{token}/{document_id}/rating.txt")

    if not os.path.isfile(path):
        data = Data(token, document_id)
        label_counter = collections.Counter(data.labels.values())
        label_counter.pop(-1)
        sentences_by_label = collections.defaultdict(
            lambda: collections.defaultdict(lambda: 0)
        )
        for item in data.get().values():
            sentence = item[0]
            label = item[1]
            sentences_by_label[label][sentence] += 1

        doc = docx.Document()
        for label, freq in label_counter.most_common():
            doc.add_paragraph(
                "Номер группы: %s     Количество предложений: %s" % (label, freq)
            )
            for sentence, freq in sentences_by_label[label].items():
                doc.add_paragraph("(%s) %s" % (freq, sentence))
            doc.add_paragraph("-" * 80)

        doc.save(path)

    return send_file(
        path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name="rating.docx",
    )


@aggregator.route("/download", methods=["GET"])
def download():
    token = session["token"]
    document_id = session["document_id"]
    struct = Struct(token, document_id).get()
    styles = Style(token, document_id).get()
    data = Data(token, document_id).sentences
    if session["toggle:reduce"] is True:
        labels = Data(token, document_id).labels
    else:
        labels = None
    content = core.compiler.assemble(
        struct,
        styles,
        data,
        labels=labels,
        mute=session["toggle:reduce"],
        out_type="txt",
        limit=40,
    )

    path = os.path.abspath(f"{UPLOAD_FOLDER}/{token}/{document_id}/result.docx")
    doc = docx.Document()
    for par in content.split("\n\n"):
        doc.add_paragraph(par)
    doc.save(path)

    with open(f"{UPLOAD_FOLDER}/{token}/{document_id}/config.json") as f:
        config = json.load(f)
    filename = config["FILENAME"]
    return send_file(
        path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name=f"{filename}.docx",
    )


@aggregator.route("/keywords", methods=["POST"])
def add_keywords():
    token = session["token"]

    for document_id in document_ids(token):
        new_keywords = {
            item.strip().casefold() for item in request.form["text"].split(",")
        }

        keywords = KeywordEntries(token, document_id).get()
        data = Data(token, document_id).sentences
        styles = Style(token, document_id).get()

        new_keyword_entries = core.search(new_keywords - keywords.keys(), data)
        styles = core.highlight(new_keyword_entries, styles)

        Style(token, document_id).post(styles)
        KeywordEntries(token, document_id).add(new_keyword_entries)

    return redirect(url_for("aggregator.document", document_id=session["document_id"]))


@aggregator.route("/keywords/<string:keyword>", methods=["GET"])
def delete_tag(keyword):
    token = session["token"]

    for document_id in document_ids(token):
        keyword_entries = {keyword: KeywordEntries(token, document_id).pop(keyword)}

        styles = Style(token, document_id).get()
        styles = core.mutelight(
            keyword_entries, styles, hide_not_matched=session[HIDE_PARS_WITHOUT_KEYWORDS]
        )

        Style(token, document_id).post(styles)

    return redirect(url_for("aggregator.document", document_id=session["document_id"]))


@aggregator.route("/hide_pars_without_keywords", methods=["GET"])
def hide_pars_without_keywords():
    token = session["token"]
    session[HIDE_PARS_WITHOUT_KEYWORDS] = not session[HIDE_PARS_WITHOUT_KEYWORDS]
    for document_id in document_ids(token):
        styles = Style(token, document_id).get()
        struct = Struct(token, document_id).get()
        keywords = KeywordEntries(token, document_id).get()
        sentensce_ids = itertools.chain(
            *((id for id, _, _ in entries) for _, entries in keywords.items())
        )
        matched_par_ids = {sentence_id.split("s")[0] for sentence_id in sentensce_ids}
        par_ids = [
            node_id
            for node_id in styles
            if core.compiler.is_paragraph_id(node_id) and node_id not in matched_par_ids
        ]
        if session[HIDE_PARS_WITHOUT_KEYWORDS]:
            styles = core.processing.hide(styles, par_ids)
        else:
            styles = core.processing.unhide(styles)

        Style(token, document_id).post(styles)
    return redirect(url_for("aggregator.document", document_id=session["document_id"]))
